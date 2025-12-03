import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from django.http import HttpResponse
from django.views import View
from django.shortcuts import render
from django import forms
from django.urls import path
from django.core.files.uploadedfile import InMemoryUploadedFile

# --- 1. The Converter Logic (Core Logic) ---

class MarkdownToDocx:
    def __init__(self):
        self.document = Document()

    def convert(self, md_text):
        """
        Converts markdown text to a docx object.
        """
        # 1. Convert Markdown to HTML using the 'tables' extension
        html = markdown.markdown(md_text, extensions=['tables', 'fenced_code'])
        
        # 2. Parse HTML to navigate the structure
        soup = BeautifulSoup(html, 'html.parser')

        # 3. Iterate through top-level elements and add to DOCX
        # Note: This is a simplified parser. For very complex nesting, 
        # recursion would be required.
        for element in soup.find_all(recursive=False):
            self._process_element(element)

        return self.document

    def _process_element(self, element):
        """Dispatches element processing based on tag name."""
        tag = element.name

        if tag.startswith('h') and len(tag) == 2:
            # Handle Headings (h1, h2, etc.)
            level = int(tag[1])
            # python-docx supports levels 0-9. 
            # We map h1->1, h2->2, etc.
            if 1 <= level <= 9:
                self.document.add_heading(element.get_text(), level=level)
        
        elif tag == 'p':
            # Handle Paragraphs
            self.document.add_paragraph(element.get_text())

        elif tag == 'ul':
            # Handle Unordered Lists
            for li in element.find_all('li'):
                self.document.add_paragraph(li.get_text(), style='List Bullet')

        elif tag == 'ol':
            # Handle Ordered Lists
            for li in element.find_all('li'):
                self.document.add_paragraph(li.get_text(), style='List Number')

        elif tag == 'table':
            # Handle Tables
            self._process_table(element)

    def _process_table(self, table_element):
        """Parses an HTML table and adds it to the DOCX."""
        rows = table_element.find_all('tr')
        if not rows:
            return

        # Determine number of columns by looking at the first row
        first_row_cols = rows[0].find_all(['th', 'td'])
        num_cols = len(first_row_cols)
        num_rows = len(rows)

        # Create table in Docx
        table = self.document.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid' # Standard grid style

        for i, row in enumerate(rows):
            cols = row.find_all(['th', 'td'])
            for j, col in enumerate(cols):
                if j < num_cols: # Prevention against malformed tables
                    cell = table.cell(i, j)
                    cell.text = col.get_text().strip()
                    
                    # Optional: Make header row bold
                    if row.find('th'):
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True

# --- 2. Django Form ---

class UploadFileForm(forms.Form):
    md_file = forms.FileField(
        label="Select a Markdown (.md) file",
        help_text="Upload a file containing markdown text, tables, and headings."
    )

# --- 3. Django Views ---

class MdToDocxView(View):
    template_name = "upload.html" # You would typically have a template file

    def get(self, request):
        # Determine if we are running as a snippet or full app
        # For this snippet, we render a simple HTML string
        form = UploadFileForm()
        return self._render_simple_ui(request, form)

    def post(self, request):
        form = UploadFileForm(request.POST, request.FILES)
        
        if form.is_valid():
            uploaded_file = request.FILES['md_file']
            
            # Read the file content
            # files are bytes, need to decode to string
            try:
                md_content = uploaded_file.read().decode('utf-8')
            except UnicodeDecodeError:
                return HttpResponse("Error: File must be UTF-8 encoded text.", status=400)

            # Perform Conversion
            converter = MarkdownToDocx()
            docx_document = converter.convert(md_content)

            # Create the HTTP Response
            response = HttpResponse(
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            filename = uploaded_file.name.replace('.md', '.docx')
            response['Content-Disposition'] = f'attachment; filename="{filename}"'

            # Save document to the response stream
            docx_document.save(response)

            return response
        
        return self._render_simple_ui(request, form)

    def _render_simple_ui(self, request, form):
        """
        Helper to render a basic HTML page without needing external templates
        for this code snippet example.
        """
        from django.template import Template, Context
        
        html_template = """
        <!DOCTYPE html>
        <html>
        <head>
            <title>Markdown to DOCX Converter</title>
            <style>
                body { font-family: sans-serif; max-width: 800px; margin: 40px auto; padding: 20px; }
                .container { border: 1px solid #ddd; padding: 20px; border-radius: 8px; }
                button { background-color: #007bff; color: white; border: none; padding: 10px 20px; cursor: pointer; }
                button:hover { background-color: #0056b3; }
                .file-input { margin-bottom: 20px; }
            </style>
        </head>
        <body>
            <div class="container">
                <h1>MD to DOCX Converter</h1>
                <p>Upload a markdown file to convert it to a Word document.</p>
                <form method="post" enctype="multipart/form-data">
                    {% csrf_token %}
                    <div class="file-input">
                        {{ form.as_p }}
                    </div>
                    <button type="submit">Convert & Download</button>
                </form>
            </div>
        </body>
        </html>
        """
        
        t = Template(html_template)
        c = Context({'form': form})
        # We need to render the csrf token manually if not using full render shortcut context
        # But render() handles this usually.
        return HttpResponse(t.render(Context({'form': form}, autoescape=False)))

# --- 4. URL Configuration (If pasting into urls.py) ---

urlpatterns = [
    path('convert/', MdToDocxView.as_view(), name='convert_md'),
]

