import json
from docx import Document
from docx.shared import Pt

def create_strategy_document(outline_json, sample_drs_text):
    """
    Takes the JSON outline and generates the full DOCX.
    """
    doc = Document()
    
    # Optional: Add a Main Title
    doc.add_heading('Test Strategy Document', 0)

    # 1. Parse the JSON response (if it's still a string)
    if isinstance(outline_json, str):
        outline_data = json.loads(outline_json)
    else:
        outline_data = outline_json

    # 2. Loop through the outline structure
    for section in outline_data:
        heading_text = section.get('title', 'Untitled')
        heading_level = section.get('level', 1)

        # --- KEY FIX: Use the 'level' directly ---
        # python-docx handles levels 1-9 natively. 
        # This solves your "retaining info of sub heading" issue.
        doc.add_heading(heading_text, level=heading_level)

        # 3. Generate content for this specific section
        # We pass the context to the LLM so it knows what to write
        section_content = generate_section_content(heading_text, sample_drs_text)
        
        # 4. Add the content to the document
        # If the content contains simple paragraphs, add them directly
        if section_content:
            doc.add_paragraph(section_content)

    doc.save('Generated_Test_Strategy.docx')
    print("Document saved successfully.")

# Mock function to represent your LLM call for content generation
def generate_section_content(heading, context):
    # This is where you would call your LLM
    # prompt = f"Write the content for the '{heading}' section based on this DRS: {context}..."
    return f"Placeholder content generated for section: {heading}..."



import re

def add_markdown_content_to_doc(doc, markdown_text):
    """
    Parses simple LLM markdown (bullets and bold) into Docx paragraphs.
    """
    lines = markdown_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Handle Bullet Points
        if line.startswith('* ') or line.startswith('- '):
            clean_text = line[2:]
            p = doc.add_paragraph(clean_text, style='List Bullet')
        
        # Handle Numbered Lists
        elif re.match(r'^\d+\.', line):
            # Remove "1. " from start
            clean_text = re.sub(r'^\d+\.\s*', '', line)
            p = doc.add_paragraph(clean_text, style='List Number')
            
        # Handle Standard Paragraphs
        else:
            p = doc.add_paragraph(line)

        # (Optional) Simple Bold handling for **text**
        # Note: This is a basic implementation. For complex markdown, 
        # consider using libraries like `markdown` or `pypandoc`.
        if "**" in line:
            # You would need a more complex parser to apply runs for bolding 
            # typically involving splitting the string by '**' and toggling bold.
            pass 


#In the main loop, replace #doc.add_paragraph(section_content) with #add_markdown_content_to_doc(doc, #section_content).